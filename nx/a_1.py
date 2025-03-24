# https://python-docx.readthedocs.io/en/latest/user/documents.html
from docx import Document
# pip install python-docx 
# f = open('wordd.docx', 'rb')
# document = Document(f)
# print(document)
source = Document('wordd.docx')    
paras = source.paragraphs


# https://ru.stackoverflow.com/questions/1237484/%D0%BA%D0%B0%D0%BA-%D0%BE%D1%82%D0%BA%D1%80%D1%8B%D1%82%D1%8C-%D1%84%D0%B0%D0%B9%D0%BB-exe-%D0%BF%D0%BE-%D0%B7%D0%B0%D0%B4%D0%B0%D0%BD%D0%BE%D0%BC%D1%83-%D0%BF%D1%83%D1%82%D0%B8-%D0%BD%D0%B0-%D0%BF%D0%B0%D0%B9%D1%82%D0%BE%D0%BD
